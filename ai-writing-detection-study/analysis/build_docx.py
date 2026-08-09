#!/usr/bin/env python3
"""Render the study's simple Markdown manuscript as a publication-style DOCX."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "paper" / "manuscript.md"
OUTPUT = ROOT / "paper" / "when-human-writing-moves.docx"


def add_runs(paragraph, text):
    for piece in re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text):
        if not piece:
            continue
        run = paragraph.add_run(piece.strip("*`") if piece[:1] in "*`" else piece)
        run.bold = piece.startswith("**")
        run.italic = piece.startswith("*") and not piece.startswith("**")
        if piece.startswith("`"):
            run.font.name = "Courier New"


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    lines = SOURCE.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[| :\-]+\|$", lines[i + 1]):
            block = [line]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i]); i += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in block]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, value in enumerate(row):
                    add_runs(table.cell(r, c).paragraphs[0], value)
                    for run in table.cell(r, c).paragraphs[0].runs:
                        run.bold = r == 0
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            p = doc.add_heading(heading.group(2), level=level if level > 1 else 0)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            add_runs(doc.add_paragraph(style="List Number"), numbered.group(1)); i += 1; continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Quote"); add_runs(p, line[2:]); i += 1; continue
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    props = doc.core_properties
    props.title = "When Human Writing Moves"
    props.subject = "AI-text detection and temporal linguistic baselines"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
